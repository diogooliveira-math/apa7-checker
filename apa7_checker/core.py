"""
apa7_checker/core.py
====================
Core logic for APA 7 reference checking.

Supports:
  - Extracting text from .docx files
  - Finding and splitting the references section
  - Validating individual reference entries against APA 7 rules
  - Classifying reference types
  - Checking alphabetical order
  - Cross-checking in-text citations vs. reference list
  - Generating HTML and JSON reports
  - Writing output files with encoding/path fallbacks
"""

from __future__ import annotations

import html
import json
import os
import re
import shutil
import sys
import tempfile
import unicodedata
import urllib.request
from datetime import datetime
from pathlib import Path
from typing import Optional

# ---------------------------------------------------------------------------
# Optional dependency: python-docx
# ---------------------------------------------------------------------------
try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False
    _DocxDocument = None  # type: ignore

# ---------------------------------------------------------------------------
# Module-level constants
# ---------------------------------------------------------------------------

_PT_MONTHS = (
    "janeiro|fevereiro|março|abril|maio|junho|"
    "julho|agosto|setembro|outubro|novembro|dezembro"
)
_EN_MONTHS = (
    "january|february|march|april|may|june|"
    "july|august|september|october|november|december"
)

# YEAR_PATTERN: matches (YYYY), (YYYYa), (YYYYb), (YYYY, Month), (YYYY, Month DD), (n.d.)
# The negative lookbehind (?<!\d) prevents matching journal issue numbers like 366(1881).
YEAR_PATTERN = re.compile(
    r'(?<!\d)\(\d{4}[a-z]?\)'
    r'|(?<!\d)\(\d{4}[a-z]?,\s*(?:' + _PT_MONTHS + r'|' + _EN_MONTHS + r')(?:\s+\d{1,2})?\)'
    r'|\(n\.d\.\)',
    re.IGNORECASE,
)

# Heading variants that mark the start of the references section
_REF_HEADING_RE = re.compile(
    r'^\s*(referências\s+bibliográficas|referências|references|bibliography|bibliografia)\s*$',
    re.IGNORECASE | re.MULTILINE,
)

# Section-end markers
_END_SECTION_RE = re.compile(
    r'^\s*(anexo|apêndice|appendix|notas|notes)\b',
    re.IGNORECASE | re.MULTILINE,
)

# Reference type classification patterns (ordered most-specific first)
_TYPE_PATTERNS = [
    (
        'conference_proceedings',
        re.compile(
            r'(In\s+[A-Z].+?\(Eds?\.\).+?(proceedings|conference|conferência|actas|anais)'
            r'|(?<!\w)(proceedings|conferência|actas|anais)\s+(of|da|do|de)\b'
            r'|(?<!\w)proceedings\b)',
            re.IGNORECASE,
        ),
    ),
    (
        'thesis',
        re.compile(r'\[(dissertação|tese|dissertation|thesis)', re.IGNORECASE),
    ),
    (
        'book_chapter',
        re.compile(r'\bIn\b.+?\(Eds?\.\)', re.IGNORECASE),
    ),
    (
        'book',
        re.compile(
            r'(?<!\w)(edition|ed\.|edi[cç][aã]o|editora|publisher|press)(?!\w)'
            r'|\..*\(Vol\.\s*\d+\)\.',  # standalone (Vol. N). after title → book series
            re.IGNORECASE,
        ),
    ),
    (
        'journal_article',
        # Match "Journal Name, 12(3)" pattern (volume+issue) or explicit vol/volume keyword.
        # But exclude entries where (Vol. N) appears after a period (book series pattern).
        re.compile(r',\s*\d+\s*\(|\bvol\b|\bvolume\b', re.IGNORECASE),
    ),
    (
        'report',
        re.compile(
            # Match known institutional/governmental author keywords, or explicit report markers.
            # Aligned with the original tool's pattern — does NOT match the generic word "report"
            # in titles, and does NOT match WHO (which is handled by website classification).
            r'(?:Gabinete|Department|Instituto|Office|Commission|Comiss[aã]o'
            r'|UNESCO|OECD|DGE|INE|European\s+Commission)'
            r'|Report\s+No\.'
            r'|Relatório\b',
            re.IGNORECASE,
        ),
    ),
    (
        'website',
        # Exclude doi.org URLs (they belong to journal/book/report entries, not websites).
        re.compile(
            r'https?://(?!doi\.org|dx\.doi\.org|data\.europa\.eu/doi|arxiv\.org)',
            re.IGNORECASE,
        ),
    ),
]

# Pattern to detect a new author token at the beginning of a line/segment
_AUTHOR_START_RE = re.compile(
    r'^[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][a-záéíóúâêîôûãõàüç\-]+,\s+[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ]\.',
)

# Pattern for lines that start a numbered list item
_NUMBERED_LINE_RE = re.compile(r'^\s*\d+[.)]\s+')


# ---------------------------------------------------------------------------
# 1. _strip_accents
# ---------------------------------------------------------------------------

def _strip_accents(s: str) -> str:
    """Remove accent marks from *s* using NFKD normalisation."""
    return ''.join(
        c for c in unicodedata.normalize('NFKD', s)
        if unicodedata.category(c) != 'Mn'
    )


# ---------------------------------------------------------------------------
# 2. load_apa_rules
# ---------------------------------------------------------------------------

def load_apa_rules(rules_path: "str | Path") -> dict:
    """Load a JSON file of APA rules.

    Returns an empty dict if the file is not found or contains invalid JSON.
    """
    try:
        p = Path(rules_path)
        with p.open(encoding='utf-8') as fh:
            return json.load(fh)
    except (FileNotFoundError, json.JSONDecodeError, TypeError, OSError):
        return {}


# ---------------------------------------------------------------------------
# 3. extract_docx_text
# ---------------------------------------------------------------------------

def extract_docx_text(docx_path: "str | Path") -> str:
    """Extract plain text from a .docx file.

    Primary strategy: copy the file to a temp directory with a safe ASCII
    filename to avoid Windows path issues with accented characters, then use
    python-docx.

    Fallback strategy: open as a ZIP archive and strip XML tags from
    word/document.xml.
    """
    docx_path = Path(docx_path)

    # --- Primary: python-docx ---
    if _DOCX_AVAILABLE:
        tmp_dir = tempfile.mkdtemp()
        try:
            safe_copy = Path(tmp_dir) / 'input.docx'
            shutil.copy2(str(docx_path), str(safe_copy))
            doc = _DocxDocument(str(safe_copy))
            paragraphs = [p.text for p in doc.paragraphs]
            return '\n'.join(paragraphs)
        except Exception:
            pass  # fall through to XML fallback
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    # --- Fallback: raw ZIP / XML ---
    try:
        import zipfile

        with zipfile.ZipFile(str(docx_path), 'r') as zf:
            with zf.open('word/document.xml') as xf:
                raw_xml = xf.read().decode('utf-8', errors='replace')
        # Strip all XML tags
        text = re.sub(r'<[^>]+>', ' ', raw_xml)
        # Collapse runs of whitespace but preserve newlines roughly
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    except Exception as exc:
        raise RuntimeError(f"Cannot extract text from {docx_path}: {exc}") from exc


# ---------------------------------------------------------------------------
# 4. find_references_section
# ---------------------------------------------------------------------------

def find_references_section(text: str) -> str:
    """Return the references section from *text*.

    Searches for a heading matching ``_REF_HEADING_RE``.  The section ends at
    the next ``_END_SECTION_RE`` marker or at the end of the document.

    If no heading is found, returns the last 25 % of the text as a heuristic
    fallback.
    """
    if not text:
        return ''

    match = _REF_HEADING_RE.search(text)
    if not match:
        # Heuristic fallback: last 25 %
        start = max(0, int(len(text) * 0.75))
        return text[start:]

    # Start AFTER the heading line so the heading itself is not fed into split_references
    section_start = match.end()
    remainder = text[section_start:]

    end_match = _END_SECTION_RE.search(remainder)
    if end_match:
        return remainder[: end_match.start()]

    return remainder


# ---------------------------------------------------------------------------
# 5. split_references
# ---------------------------------------------------------------------------

def split_references(refs_text: str) -> list:
    """Split a block of reference text into individual reference strings.

    Strategy:
      1. Split on newlines.
      2. Strip leading list numbers (``1.``, ``1)``, etc.).
      3. Identify entry-start lines (capital surname + initial, or
         institutional/all-caps author + year).
      4. Append continuation lines to the current entry.
    """
    if not refs_text or not refs_text.strip():
        return []

    lines = refs_text.splitlines()

    # Institutional author pattern: starts with capital word(s) followed by year
    _INST_START_RE = re.compile(
        r'^[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇa-záéíóúâêîôûãõàüç\s&,\.\-]+'
        r'\s+\(\d{4}',
    )

    def _is_entry_start(line: str) -> bool:
        stripped = line.strip()
        if not stripped:
            return False
        # Remove leading numbering
        stripped = _NUMBERED_LINE_RE.sub('', stripped)
        # Personal author: Surname, I.
        if _AUTHOR_START_RE.match(stripped):
            return True
        # Institutional / org author followed by year
        if _INST_START_RE.match(stripped):
            return True
        return False

    cleaned_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        # Remove leading numbering
        stripped = _NUMBERED_LINE_RE.sub('', stripped).strip()
        if stripped:
            cleaned_lines.append(stripped)

    if not cleaned_lines:
        return []

    entries: list[str] = []
    current_parts: list[str] = []

    for line in cleaned_lines:
        if _is_entry_start(line):
            if current_parts:
                entries.append(' '.join(current_parts))
            current_parts = [line]
        else:
            if current_parts:
                current_parts.append(line)
            else:
                # No entry started yet – treat as a new one
                current_parts = [line]

    if current_parts:
        entries.append(' '.join(current_parts))

    return [e.strip() for e in entries if e.strip()]


# ---------------------------------------------------------------------------
# 6. is_personal_communication
# ---------------------------------------------------------------------------

def is_personal_communication(entry: str) -> bool:
    """Return True if *entry* describes a personal communication."""
    return bool(re.search(r'comunicação\s+pessoal|personal\s+communication', entry, re.IGNORECASE))


# ---------------------------------------------------------------------------
# 7. has_doi
# ---------------------------------------------------------------------------

def has_doi(entry: str) -> bool:
    """Return True if *entry* contains a DOI or URL."""
    return bool(re.search(r'https?://|doi\.org|doi:', entry, re.IGNORECASE))


# ---------------------------------------------------------------------------
# 8. validate_reference_entry
# ---------------------------------------------------------------------------

def validate_reference_entry(entry: str, rules: Optional[dict] = None) -> dict:
    """Heuristic APA 7 checks for a single reference entry.

    Returns::

        {
            "entry": str,
            "issues": [str, ...],
            "suggestion": str,
        }
    """
    issues: list[str] = []

    if not entry or not entry.strip():
        return {"entry": entry, "issues": ["Empty entry"], "suggestion": ""}

    # ------------------------------------------------------------------
    # Author format check
    # ------------------------------------------------------------------
    if not is_personal_communication(entry):
        # Extract the part before the year parenthesis
        author_part = entry.split('(')[0].strip() if '(' in entry else entry.strip()
        has_author_comma = ',' in author_part

        # Institutional / organisational authors (e.g. "World Health Organization.",
        # "OECD.", "UNESCO.", "Vodafone Foundation.", "Comissão Europeia.") are valid
        # without commas in APA 7.
        # Heuristic: treat as institutional if:
        #   (a) single all-caps abbreviation (OECD, WHO), OR
        #   (b) 3+ words allowing lowercase connectors, first word capitalised, no
        #       "Surname, I." or initials patterns (same threshold as original tool), OR
        #   (c) 2 words where the LAST word is a known org suffix
        words_in_author = [w for w in author_part.rstrip('.').split() if w]
        _CONNECTORS = {
            'de', 'do', 'da', 'dos', 'das', 'di', 'del', 'der',
            'for', 'of', 'the', 'and', 'e', 'y', 'und', 'en', 'et',
            'a', 'o', 'os', 'as', 'para', 'per', 'por', 'van', 'von',
        }
        # Known org suffixes that reliably indicate an organisation name
        _ORG_SUFFIXES = {
            'organization', 'organisation', 'foundation', 'commission', 'committee',
            'association', 'institute', 'institution', 'agency', 'authority',
            'council', 'department', 'office', 'bureau', 'ministry', 'secretariat',
            'federation', 'union', 'alliance', 'network', 'coalition', 'programme',
            'program', 'centre', 'center', 'laboratory', 'group', 'division',
            'directorate', 'presidência', 'ministério', 'departamento', 'instituto',
            'fundação', 'comissão', 'agência', 'conselho', 'assembleia',
        }
        is_abbreviation = bool(
            len(words_in_author) == 1 and words_in_author[0].isupper()
        )
        # 3+ word org (original threshold, with lowercase connectors allowed)
        _is_3plus_word_org = (
            len(words_in_author) >= 3
            and words_in_author[0][0].isupper()
            and all(w[0].isupper() or w.lower() in _CONNECTORS for w in words_in_author)
            and not re.search(r',\s*[A-Z]\.\s*[A-Z]\.', author_part)
            and not re.search(r',\s*[A-Z]\.', author_part)
        )
        # 2-word org where either the first or last word is a known org suffix
        _is_2word_org_with_suffix = (
            len(words_in_author) == 2
            and words_in_author[0][0].isupper()
            and (
                words_in_author[1].lower().rstrip('.') in _ORG_SUFFIXES
                or words_in_author[0].lower().rstrip('.') in _ORG_SUFFIXES
            )
        )
        is_multi_word_org = _is_3plus_word_org or _is_2word_org_with_suffix
        is_institutional = is_abbreviation or is_multi_word_org

        if not has_author_comma and not is_institutional:
            issues.append(
                "Author format: no comma found before year — expected 'Surname, I.' format"
            )
        elif has_author_comma:
            # Check that first token looks like a surname (starts uppercase)
            first_word = author_part.strip().split()[0] if author_part.strip() else ''
            if first_word and first_word[0].islower():
                issues.append(
                    "Author format: entry appears to start with a lowercase word — "
                    "use 'Surname, I.' format"
                )

    # ------------------------------------------------------------------
    # Year check
    # ------------------------------------------------------------------
    year_match = YEAR_PATTERN.search(entry)
    if not year_match:
        issues.append("Year not found or invalid format (expected e.g. (2023) or (n.d.))")

    # ------------------------------------------------------------------
    # Semicolon-separated authors check
    # ------------------------------------------------------------------
    # APA 7 uses "&" and commas, never semicolons between authors.
    author_section = entry.split('(')[0] if '(' in entry else entry
    if ';' in author_section:
        issues.append(
            "Author format: semicolons found between authors — "
            "APA 7 uses commas and '&' (e.g. 'Silva, J., & Ramos, P.')"
        )

    # ------------------------------------------------------------------
    # DOI format check
    # ------------------------------------------------------------------
    # Bad DOI prefixes
    bad_doi_patterns = [
        r'http://doi\.org/',
        r'http://dx\.doi\.org/',
        r'https?://dx\.doi\.org/',
        r'(?<![:/])doi:\s*10\.',
        r'(?<![:/\w])doi\.org/(?!.*https?://)',
    ]
    for pat in bad_doi_patterns:
        if re.search(pat, entry, re.IGNORECASE):
            issues.append(
                "DOI format: use 'https://doi.org/...' (not http://, dx.doi.org, or bare doi:)"
            )
            break

    # ------------------------------------------------------------------
    # Publisher location (APA 7 removed city requirement)
    # ------------------------------------------------------------------
    # Only flag the explicit "City, ST: Publisher" pattern (e.g. "London, UK: Routledge").
    # The original tool used: r'\b[A-Z][a-z]+,\s*[A-Z]{2,3}:\s*[A-Za-z]'
    # We use the same pattern to avoid false positives from subtitle colons or URLs.
    if re.search(r'\b[A-Z][a-z]+,\s*[A-Z]{2,3}:\s*[A-Za-z]', entry):
        issues.append(
            "Publisher location: APA 7 no longer requires city/location before publisher name"
        )

    # ------------------------------------------------------------------
    # Truncation
    # ------------------------------------------------------------------
    stripped = entry.strip()
    if stripped.endswith('...') or re.search(r'\bet\s+al\.\s*$', stripped):
        issues.append("Entry appears truncated (ends with '...' or 'et al.')")

    suggestion = ""
    if issues:
        suggestion = (
            "Review this entry against APA 7 guidelines. "
            "Key points: Surname, I. (Year). Title. Source. https://doi.org/xxx"
        )

    return {"entry": entry, "issues": issues, "suggestion": suggestion}


# ---------------------------------------------------------------------------
# 9. validate_year_pattern
# ---------------------------------------------------------------------------

def validate_year_pattern(entry: str) -> dict:
    """Check only the year expression in *entry*.

    Returns::

        {"valid": bool, "year_found": str | None, "issues": [str, ...]}
    """
    issues: list[str] = []
    match = YEAR_PATTERN.search(entry)
    if match:
        return {"valid": True, "year_found": match.group(), "issues": issues}

    issues.append("Year not found or invalid format (expected e.g. (2023) or (n.d.))")
    return {"valid": False, "year_found": None, "issues": issues}


# ---------------------------------------------------------------------------
# 10. detect_jammed_entries
# ---------------------------------------------------------------------------

def detect_jammed_entries(entry: str) -> "list[str] | None":
    """Detect and attempt to split a jammed (merged) reference entry.

    If more than one YEAR_PATTERN match is found, try to split at a new-author
    boundary.  Returns a list of parts if successful, else None.
    """
    matches = list(YEAR_PATTERN.finditer(entry))
    if len(matches) <= 1:
        return None

    # If all year matches share the same value (e.g. a title like
    # "(COM(2018) 22 final; SWD(2018) 12 final)" repeats the same year),
    # the entry is almost certainly NOT jammed — it is a single reference
    # with the same year appearing multiple times.
    distinct_years = {m.group() for m in matches}
    if len(distinct_years) <= 1:
        return None

    # Phase 1: split at ". " before an apparent new author token
    # Look for ". Surname, I." pattern
    split_re = re.compile(
        r'\.\s+(?=[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][a-záéíóúâêîôûãõàüç\-]+,\s+[A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ]\.)'
    )
    parts_phase1 = split_re.split(entry)
    if len(parts_phase1) >= 2:
        # Restore the period that was consumed
        restored = []
        for i, part in enumerate(parts_phase1):
            if i < len(parts_phase1) - 1:
                restored.append(part.rstrip() + '.')
            else:
                restored.append(part)
        cleaned = [p.strip() for p in restored if p.strip()]
        if len(cleaned) >= 2:
            return cleaned

    # Phase 2: fallback — split at the position just before the second year match
    second_match = matches[1]
    # Find the last ". " before the second year match
    segment_before = entry[: second_match.start()]
    last_period = segment_before.rfind('. ')
    if last_period != -1:
        part1 = entry[: last_period + 1].strip()
        part2 = entry[last_period + 2:].strip()
        if part1 and part2:
            return [part1, part2]

    # Final fallback: split exactly at the second match's start
    part1 = entry[: second_match.start()].strip()
    part2 = entry[second_match.start():].strip()
    if part1 and part2:
        return [part1, part2]

    return None


# ---------------------------------------------------------------------------
# 11. extract_docx_refs_with_italic
# ---------------------------------------------------------------------------

def extract_docx_refs_with_italic(docx_path: "str | Path") -> dict:
    """Extract run-level italic formatting for each paragraph in the
    references section.

    Returns::

        {
            normalized_text: {
                "text": str,
                "italic_spans": [[start, end], ...],
            },
            ...
        }

    If python-docx is not available, returns an empty dict.
    """
    if not _DOCX_AVAILABLE:
        return {}

    docx_path = Path(docx_path)
    result: dict = {}

    tmp_dir = tempfile.mkdtemp()
    try:
        safe_copy = Path(tmp_dir) / 'input.docx'
        shutil.copy2(str(docx_path), str(safe_copy))
        doc = _DocxDocument(str(safe_copy))
    except Exception:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return {}
    finally:
        pass  # cleanup happens after processing

    try:
        full_text = '\n'.join(p.text for p in doc.paragraphs)
        ref_section_text = find_references_section(full_text)
        ref_lines = set(ref_section_text.splitlines())

        for para in doc.paragraphs:
            para_text = para.text
            if para_text.strip() not in ref_lines and para_text.strip() == '':
                continue

            italic_spans: list[list[int]] = []
            offset = 0
            for run in para.runs:
                run_len = len(run.text)
                if run.italic:
                    italic_spans.append([offset, offset + run_len])
                offset += run_len

            # Merge adjacent spans
            merged: list[list[int]] = []
            for span in italic_spans:
                if merged and span[0] <= merged[-1][1]:
                    merged[-1][1] = max(merged[-1][1], span[1])
                else:
                    merged.append(span[:])

            key = para_text.strip()
            result[key] = {"text": para_text, "italic_spans": merged}
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    return result


# ---------------------------------------------------------------------------
# 12. _italic_text
# ---------------------------------------------------------------------------

def _italic_text(entry_text: str, italic_spans: list) -> str:
    """Return the concatenation of all italic substrings."""
    if not italic_spans:
        return ''
    parts: list[str] = []
    for span in italic_spans:
        start, end = span[0], span[1]
        parts.append(entry_text[start:end])
    return ''.join(parts)


# ---------------------------------------------------------------------------
# 13. _non_italic_text
# ---------------------------------------------------------------------------

def _non_italic_text(entry_text: str, italic_spans: list) -> str:
    """Return the concatenation of all non-italic substrings."""
    if not italic_spans:
        return entry_text
    parts: list[str] = []
    prev = 0
    for span in italic_spans:
        start, end = span[0], span[1]
        if prev < start:
            parts.append(entry_text[prev:start])
        prev = end
    if prev < len(entry_text):
        parts.append(entry_text[prev:])
    return ''.join(parts)


# ---------------------------------------------------------------------------
# 14. validate_italic_formatting
# ---------------------------------------------------------------------------

def validate_italic_formatting(
    entry_text: str,
    italic_spans: list,
    ref_type: str,
) -> list:
    """Return a list of italic-formatting issue strings (empty = no issues).

    Rules are applied per reference type.  If *italic_spans* is empty/None,
    the check is skipped.
    """
    if not italic_spans:
        return []
    if not entry_text:
        return []

    issues: list[str] = []
    italic = _italic_text(entry_text, italic_spans)
    non_italic = _non_italic_text(entry_text, italic_spans)

    if ref_type == 'journal_article':
        # Journal name + volume should be italic; article title should NOT be italic
        # Heuristic: after the year "(YYYY)." there's a title, then journal name
        year_m = YEAR_PATTERN.search(entry_text)
        if year_m:
            after_year = entry_text[year_m.end():].strip().lstrip('.')
            # First sentence-like chunk is the article title
            title_end = after_year.find('.')
            if title_end != -1:
                article_title = after_year[:title_end].strip()
                # Check article title is NOT italic
                if article_title and article_title in italic:
                    issues.append(
                        "Journal article: article title should NOT be italic (only journal name and volume)"
                    )
        # Volume pattern e.g. 23(4) — the issue number in parens should not be italic
        vol_match = re.search(r'(\d+)\((\d+)\)', entry_text)
        if vol_match:
            issue_num = f'({vol_match.group(2)})'
            if issue_num in italic:
                issues.append(
                    "Journal article: issue number in parentheses should NOT be italic"
                )
            vol_num = vol_match.group(1)
            # Volume number itself should be italic — check by span overlap
            vol_start = vol_match.start(1)
            vol_end = vol_match.end(1)
            in_italic = any(
                s[0] <= vol_start and vol_end <= s[1] for s in italic_spans
            )
            if not in_italic:
                issues.append(
                    "Journal article: journal name and volume number should be italic"
                )

    elif ref_type == 'book':
        # Main title should be italic — check that some italic text exists after year
        year_m = YEAR_PATTERN.search(entry_text)
        if year_m:
            after_year = entry_text[year_m.end():]
            if not any(
                s[0] >= year_m.end() for s in italic_spans
            ):
                issues.append("Book: main title should be italic")

    elif ref_type == 'book_chapter':
        # Chapter title NOT italic; book title (after "In") SHOULD be italic
        in_match = re.search(r'\bIn\b', entry_text)
        if in_match:
            before_in = entry_text[: in_match.start()]
            # Check chapter title (between year and "In") is not italic
            year_m = YEAR_PATTERN.search(before_in)
            if year_m:
                chapter_segment = before_in[year_m.end():]
                if chapter_segment.strip() and chapter_segment.strip() in italic:
                    issues.append(
                        "Book chapter: chapter title should NOT be italic"
                    )
            # Book title after "In editor (Ed.)" should be italic
            after_in = entry_text[in_match.end():]
            eds_match = re.search(r'\(Eds?\.\)', after_in)
            # Book title after "In editor (Ed.)" should be italic.
            # The book title ends at the first " (pp.", " (Vol.", "(1st ed.", "(" or "."
            # We must truncate there to avoid including page-range/edition info.
            if eds_match:
                book_title_seg = after_in[eds_match.end():].strip().lstrip(',').strip()
                # Stop at opening parenthesis (edition, page range, volume) or standalone "."
                paren_match = re.search(r'\s*\(', book_title_seg)
                first_dot = book_title_seg.find('.')
                # Pick whichever comes first
                stops = [pos for pos in [
                    paren_match.start() if paren_match else None,
                    first_dot if first_dot != -1 else None,
                ] if pos is not None]
                if stops:
                    book_title = book_title_seg[:min(stops)].strip()
                else:
                    book_title = book_title_seg.strip()
                if book_title and book_title not in italic:
                    issues.append(
                        "Book chapter: book title (after editor) should be italic"
                    )

    elif ref_type == 'thesis':
        # Title in brackets should be italic
        bracket_match = re.search(r'\[([^\]]+)\]', entry_text)
        if bracket_match:
            bracket_content = bracket_match.group(1)
            b_start = bracket_match.start(1)
            b_end = bracket_match.end(1)
            in_italic = any(
                s[0] <= b_start and b_end <= s[1] for s in italic_spans
            )
            if not in_italic:
                issues.append("Thesis: title in brackets should be italic")

    elif ref_type in ('report', 'website'):
        # Main title should be italic
        year_m = YEAR_PATTERN.search(entry_text)
        if year_m:
            if not any(s[0] >= year_m.end() for s in italic_spans):
                issues.append(f"{ref_type.capitalize()}: main title should be italic")

    return issues


# ---------------------------------------------------------------------------
# 15. classify_reference_type
# ---------------------------------------------------------------------------

def classify_reference_type(entry: str, rules: Optional[dict] = None) -> str:
    """Classify the reference type by iterating ``_TYPE_PATTERNS``."""
    for type_name, pattern in _TYPE_PATTERNS:
        if pattern.search(entry):
            return type_name
    return 'other'


# ---------------------------------------------------------------------------
# 16. validate_by_type
# ---------------------------------------------------------------------------

def validate_by_type(entry: str, ref_type: str) -> dict:
    """Type-specific APA 7 checks.

    Returns::

        {"type": ref_type, "type_issues": [str, ...], "suggestion": str}
    """
    type_issues: list[str] = []

    if ref_type == 'journal_article':
        # Must have volume(issue), pp–pp pattern  (original rule: page range only)
        if not re.search(r',\s*\d{1,4}(?:\(\d{1,4}\))?,\s*\d+\s*[–\-]\s*\d+', entry):
            type_issues.append(
                "Journal article: missing page range — expected volume(issue), pp–pp format"
            )
        # NOTE: DOI is encouraged by APA7 but many older references don't have one.
        # The original validator did NOT flag missing DOI as an error, so we don't either.

    elif ref_type == 'book':
        # Should NOT have journal-style volume(issue) numbers
        if re.search(r',\s*\d+\s*\(\d+\)', entry):
            type_issues.append(
                "Book: entry contains journal-style volume(issue) — check reference type"
            )

    elif ref_type == 'book_chapter':
        # Must have pp. xx–xx
        if not re.search(r'pp\.\s*\d+\s*[–\-]\s*\d+', entry):
            type_issues.append(
                "Book chapter: missing page range — expected (pp. xx–xx) in chapter entry"
            )
        # Must have "In Initials Surname (Ed./Eds.)"
        if not re.search(r'\bIn\s+[A-Z]', entry):
            type_issues.append(
                "Book chapter: missing editor info — expected \"In Initials Surname (Ed[s].)\""
            )
        # NOTE: italic validation (book title should be italic) is handled separately
        # by validate_italic_formatting() using real run-level data, not text heuristics.

    elif ref_type == 'conference_proceedings':
        # Must have a bracket description
        if not re.search(r'\[[^\]]+\]', entry):
            type_issues.append(
                "Conference paper: missing bracket label — expected [Paper presentation] or similar"
            )

    elif ref_type == 'thesis':
        # Must have a bracket description or a repository URL
        if not re.search(r'\[[^\]]+\]', entry) and 'http' not in entry.lower():
            type_issues.append(
                "Thesis: missing bracket description or URL — expected "
                "[Unpublished master's thesis, Institution] or repository URL"
            )

    elif ref_type == 'report':
        # Should have a URL or DOI (original rule)
        if not re.search(r'https?://', entry):
            type_issues.append(
                "Report: missing URL — institutional reports should include a URL or DOI"
            )

    elif ref_type == 'website':
        # URL already implied by classification, but double-check
        if not re.search(r'https?://', entry):
            type_issues.append("Website: missing URL")

    suggestion = ""
    if type_issues:
        suggestion = f"Review {ref_type.replace('_', ' ')} formatting per APA 7."

    return {"type": ref_type, "type_issues": type_issues, "suggestion": suggestion}


# ---------------------------------------------------------------------------
# 17. _extract_intext_keys
# ---------------------------------------------------------------------------

def _extract_intext_keys(body_text: str) -> set:
    """Extract all in-text citation keys as (surname, year) tuples.

    Handles:
      - ``(Author, YYYY)``
      - ``(Author & Author, YYYY)``
      - ``(Author et al., YYYY)``
      - ``Author (YYYY)``

    Year keys are normalised by stripping trailing letter suffixes (a, b, …)
    so that ``Wing (2008)`` in the body matches ``Wing (2008a)`` in the
    reference list.
    """
    keys: set = set()

    def _norm_key(s: str) -> str:
        return _strip_accents(s.strip().lower())

    def _norm_year(raw: str) -> str:
        """Normalise year: strip parens/spaces, then strip trailing letter."""
        y = re.sub(r'[\(\)\s]', '', raw).lower()
        # Strip trailing letter suffix: 2008a → 2008
        return re.sub(r'^(\d{4})[a-z]$', r'\1', y)

    # Pattern 1: (Surname, YYYY) or (Surname & Surname, YYYY) or (Surname et al., YYYY)
    paren_pattern = re.compile(
        r'\(([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][^\d\(\)]{1,80}?),\s*(\d{4}[a-z]?|n\.d\.)\)',
        re.IGNORECASE,
    )
    for m in paren_pattern.finditer(body_text):
        authors_raw = m.group(1).strip()
        year = _norm_year(m.group(2))
        # Take the first surname
        first_author = re.split(r'\s*&\s*|\s*,\s*|\s+et\s+al', authors_raw)[0].strip()
        if first_author:
            keys.add((_norm_key(first_author), year))

    # Pattern 2: Surname (YYYY) narrative citation
    narrative_pattern = re.compile(
        r'([A-ZÁÉÍÓÚÂÊÎÔÛÃÕÀÜÇ][a-záéíóúâêîôûãõàüç\-]+)\s+\((\d{4}[a-z]?|n\.d\.)\)',
    )
    for m in narrative_pattern.finditer(body_text):
        surname = _norm_key(m.group(1))
        year = _norm_year(m.group(2))
        keys.add((surname, year))

    return keys


# ---------------------------------------------------------------------------
# 18. _extract_ref_key
# ---------------------------------------------------------------------------

def _extract_ref_key(entry: str) -> "tuple | None":
    """Extract (surname, year) from a reference entry.

    Returns None if unable to parse.
    Year letter suffixes (2008a → 2008) are stripped so that reference list
    keys match in-text citation keys.
    """
    if not entry or not entry.strip():
        return None

    def _norm_key(s: str) -> str:
        return _strip_accents(s.strip().lower())

    def _norm_year(raw: str) -> str:
        y = re.sub(r'[\(\)\s]', '', raw).lower()
        return re.sub(r'^(\d{4})[a-z]$', r'\1', y)

    # Find year
    year_match = YEAR_PATTERN.search(entry)
    if not year_match:
        return None
    year = _norm_year(year_match.group())

    # Find surname: first word before the comma
    first_comma = entry.find(',')
    if first_comma == -1:
        # Institutional author: use first word
        first_word = entry.strip().split()[0]
        return (_norm_key(first_word), year)

    surname = entry[:first_comma].strip()
    return (_norm_key(surname), year)


# ---------------------------------------------------------------------------
# 19. cross_check_citations
# ---------------------------------------------------------------------------

def cross_check_citations(body_text: str, reference_entries: list) -> dict:
    """Cross-check in-text citations against the reference list.

    Returns::

        {
            "cited_not_listed": [(surname, year), ...],
            "listed_not_cited": [(surname, year), ...],
            "matched": [(surname, year), ...],
        }
    """
    cited = _extract_intext_keys(body_text)
    listed: set = set()
    for entry in reference_entries:
        key = _extract_ref_key(entry)
        if key:
            listed.add(key)

    matched = cited & listed
    cited_not_listed = sorted(cited - listed)
    listed_not_cited = sorted(listed - cited)

    return {
        "cited_not_listed": cited_not_listed,
        "listed_not_cited": listed_not_cited,
        "matched": sorted(matched),
    }


# ---------------------------------------------------------------------------
# 20. validate_doi_live
# ---------------------------------------------------------------------------

def validate_doi_live(entry_or_url: str, timeout: int = 5) -> dict:
    """Validate a DOI or URL by performing an HTTP HEAD request.

    Returns::

        {"url": str, "status": int | None, "reachable": bool, "error": str | None}
    """
    # Extract URL from entry if not a direct URL
    url: Optional[str] = None
    if re.match(r'https?://', entry_or_url.strip(), re.IGNORECASE):
        url = entry_or_url.strip()
    else:
        url_match = re.search(r'https?://\S+', entry_or_url)
        if url_match:
            url = url_match.group().rstrip('.,;)')
        else:
            doi_match = re.search(r'doi:\s*(10\.\S+)', entry_or_url, re.IGNORECASE)
            if doi_match:
                url = f"https://doi.org/{doi_match.group(1).rstrip('.,;)')}"

    if not url:
        return {"url": "", "status": None, "reachable": False, "error": "No URL or DOI found in entry"}

    try:
        req = urllib.request.Request(url, method='HEAD')
        req.add_header('User-Agent', 'apa7-checker/1.0 (reference validator)')
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            status = resp.status
            return {"url": url, "status": status, "reachable": status < 400, "error": None}
    except urllib.error.HTTPError as exc:
        return {"url": url, "status": exc.code, "reachable": exc.code < 400, "error": str(exc)}
    except urllib.error.URLError as exc:
        return {"url": url, "status": None, "reachable": False, "error": str(exc.reason)}
    except OSError as exc:
        return {"url": url, "status": None, "reachable": False, "error": str(exc)}
    except Exception as exc:  # noqa: BLE001
        return {"url": url, "status": None, "reachable": False, "error": str(exc)}


# ---------------------------------------------------------------------------
# 21. _sort_key_for_entry
# ---------------------------------------------------------------------------

def _sort_key_for_entry(entry: str) -> tuple:
    """Return ``(normalized_surname, normalized_initials, normalized_year)``
    for sorting purposes.

    All components are accent-stripped, lowercased, non-alphanumeric removed.
    """
    def _clean(s: str) -> str:
        return re.sub(r'[^a-z0-9]', '', _strip_accents(s).lower())

    if not entry or not entry.strip():
        return ('', '', '')

    year_match = YEAR_PATTERN.search(entry)
    year_raw = year_match.group() if year_match else ''
    year_clean = _clean(year_raw)

    first_comma = entry.find(',')
    # A comma only belongs to the author field if it appears BEFORE the year.
    # If the first comma is after (or there is no year), treat as institutional.
    if first_comma == -1 or (year_match and first_comma > year_match.start()):
        # Institutional author — use full pre-year text
        pre_year = entry[: year_match.start()].strip() if year_match else entry
        return (_clean(pre_year), '', year_clean)

    surname = entry[:first_comma].strip()
    after_comma = entry[first_comma + 1:]

    # Initials: up to the year
    if year_match:
        initials_raw = entry[first_comma + 1: year_match.start()]
    else:
        # Take first two "words"
        initials_raw = ' '.join(after_comma.split()[:2])

    return (_clean(surname), _clean(initials_raw), year_clean)


# ---------------------------------------------------------------------------
# 22. check_alphabetical_order
# ---------------------------------------------------------------------------

def check_alphabetical_order(entries: list) -> dict:
    """Check if reference entries are in APA 7 alphabetical order.

    Returns::

        {
            "is_sorted": bool,
            "out_of_order": [str, ...],
            "correct_order": [str, ...],
            "word_instructions": str,
        }
    """
    if not entries:
        return {
            "is_sorted": True,
            "out_of_order": [],
            "correct_order": [],
            "word_instructions": "",
        }

    sorted_entries = sorted(entries, key=_sort_key_for_entry)
    is_sorted = entries == sorted_entries

    out_of_order: list[str] = []
    if not is_sorted:
        for orig, correct in zip(entries, sorted_entries):
            if orig != correct:
                out_of_order.append(orig)

    word_instructions = generate_word_instructions(entries, sorted_entries, out_of_order)

    return {
        "is_sorted": is_sorted,
        "out_of_order": out_of_order,
        "correct_order": sorted_entries,
        "word_instructions": word_instructions,
    }


# ---------------------------------------------------------------------------
# 23. generate_word_instructions
# ---------------------------------------------------------------------------

def generate_word_instructions(
    original: list,
    sorted_: list,
    out_of_order: list,
) -> str:
    """Generate plain-text Word reordering instructions.

    Entries that need to move are marked with ``►``.
    """
    if not original:
        return ""

    lines: list[str] = [
        "WORD REORDERING INSTRUCTIONS",
        "=" * 60,
        "Paste the references in the order listed below.",
        "Entries marked with ► are currently out of order.",
        "",
    ]

    out_of_order_set = set(out_of_order)
    for i, entry in enumerate(sorted_, start=1):
        marker = "► " if entry in out_of_order_set else "  "
        # Truncate long entries for readability
        display = entry if len(entry) <= 120 else entry[:117] + "..."
        lines.append(f"{marker}{i:3d}. {display}")

    if not out_of_order:
        lines.append("")
        lines.append("✓ References are already in correct alphabetical order.")

    return '\n'.join(lines)


# ---------------------------------------------------------------------------
# 24. _norm
# ---------------------------------------------------------------------------

def _norm(s: str) -> str:
    """Normalise whitespace: strip and collapse multiple spaces."""
    return re.sub(r'\s+', ' ', s).strip()


# ---------------------------------------------------------------------------
# 25. generate_html_report
# ---------------------------------------------------------------------------

_HTML_CSS = """
* { box-sizing: border-box; margin: 0; padding: 0; }
body { font-family: 'Segoe UI', Arial, sans-serif; background: #f4f6f9; color: #222; line-height: 1.5; }
.container { max-width: 1200px; margin: 0 auto; padding: 2rem; }
h1 { font-size: 1.8rem; color: #1a237e; margin-bottom: 0.25rem; }
h2 { font-size: 1.25rem; color: #283593; margin: 2rem 0 0.75rem; border-bottom: 2px solid #c5cae9; padding-bottom: 0.25rem; }
p.meta { font-size: 0.85rem; color: #555; margin-bottom: 0.5rem; }

/* Summary grid */
.summary-grid { display: grid; grid-template-columns: repeat(auto-fill, minmax(160px, 1fr)); gap: 1rem; margin-bottom: 1.5rem; }
.card { background: #fff; border-radius: 8px; padding: 1rem; text-align: center; box-shadow: 0 1px 4px rgba(0,0,0,.1); }
.card .card-value { font-size: 2rem; font-weight: 700; color: #1a237e; }
.card .card-label { font-size: 0.8rem; color: #666; margin-top: 0.2rem; }
.card.warn .card-value { color: #e65100; }
.card.error .card-value { color: #b71c1c; }
.card.ok .card-value { color: #1b5e20; }

/* Badges */
.badge { display: inline-block; padding: 0.2em 0.6em; border-radius: 12px; font-size: 0.75rem; font-weight: 600; }
.badge-ok { background: #e8f5e9; color: #1b5e20; }
.badge-warn { background: #fff3e0; color: #e65100; }
.badge-error { background: #ffebee; color: #b71c1c; }
.type-tag { background: #e3f2fd; color: #0d47a1; padding: 0.15em 0.5em; border-radius: 10px; font-size: 0.75rem; }
.alpha-badge-ok { background: #e8f5e9; color: #1b5e20; padding: 0.15em 0.5em; border-radius: 10px; font-size: 0.75rem; }
.alpha-badge-bad { background: #ffebee; color: #b71c1c; padding: 0.15em 0.5em; border-radius: 10px; font-size: 0.75rem; }

/* Table */
table { width: 100%; border-collapse: collapse; background: #fff; border-radius: 8px; overflow: hidden; box-shadow: 0 1px 4px rgba(0,0,0,.1); font-size: 0.88rem; }
thead { background: #1a237e; color: #fff; }
th { padding: 0.6rem 0.75rem; text-align: left; font-weight: 600; }
tbody tr:nth-child(even) { background: #f5f5f5; }
tbody tr:hover { background: #e8eaf6; }
td { padding: 0.55rem 0.75rem; vertical-align: top; border-bottom: 1px solid #e0e0e0; }

/* Issue list */
.issue-list { list-style: disc; padding-left: 1.2rem; color: #b71c1c; font-size: 0.82rem; }
.issue-list li { margin-bottom: 0.2rem; }
.sugg { font-size: 0.82rem; color: #1a237e; font-style: italic; margin-top: 0.3rem; }
.orig { font-size: 0.82rem; color: #666; font-style: italic; max-width: 400px; }

/* Alpha section */
.alpha-ok { color: #1b5e20; font-weight: 600; }
.alpha-bad { color: #b71c1c; font-weight: 600; }
.alpha-list { list-style: decimal; padding-left: 1.5rem; margin-top: 0.5rem; font-size: 0.85rem; }
.alpha-list li { margin-bottom: 0.2rem; }
"""


def generate_html_report(report: dict, out_path: "str | Path | None" = None) -> str:
    """Render a complete self-contained HTML report.

    If *out_path* is provided, write the HTML to file via ``safe_write_file``.
    Always returns the HTML string.
    """
    source_file = html.escape(str(report.get("source_file", "unknown")))
    timestamp = html.escape(report.get("timestamp", datetime.now().isoformat(timespec='seconds')))

    entries_data: list[dict] = report.get("entries", [])
    alpha_data: dict = report.get("alphabetical_order", {})

    total = len(entries_data)
    ok_count = sum(1 for e in entries_data if not e.get("all_issues"))
    issues_count = total - ok_count
    manual_count = sum(1 for e in entries_data if e.get("needs_manual_lookup"))
    jammed_count = sum(1 for e in entries_data if e.get("jammed_split"))
    order_errors = len(alpha_data.get("out_of_order", []))

    # ---------- Summary cards ----------
    cards_html = f"""
<div class="summary-grid">
  <div class="card"><div class="card-value">{total}</div><div class="card-label">References</div></div>
  <div class="card ok"><div class="card-value">{ok_count}</div><div class="card-label">OK</div></div>
  <div class="card {'error' if issues_count else 'ok'}"><div class="card-value">{issues_count}</div><div class="card-label">With Issues</div></div>
  <div class="card {'warn' if manual_count else 'ok'}"><div class="card-value">{manual_count}</div><div class="card-label">Manual Check</div></div>
  <div class="card {'warn' if jammed_count else 'ok'}"><div class="card-value">{jammed_count}</div><div class="card-label">Jammed</div></div>
  <div class="card {'error' if order_errors else 'ok'}"><div class="card-value">{order_errors}</div><div class="card-label">Order Errors</div></div>
</div>
"""

    # ---------- Alpha section ----------
    is_sorted = alpha_data.get("is_sorted", True)
    if is_sorted:
        alpha_html = '<p class="alpha-ok">✓ References are in correct alphabetical order.</p>'
    else:
        ooo = alpha_data.get("out_of_order", [])
        items = '\n'.join(f'<li>{html.escape(e)}</li>' for e in ooo)
        alpha_html = (
            f'<p class="alpha-bad">✗ {len(ooo)} reference(s) are out of order.</p>'
            f'<ul class="alpha-list">{items}</ul>'
        )

    # ---------- Table rows ----------
    rows: list[str] = []
    for idx, ed in enumerate(entries_data, start=1):
        entry_text = html.escape(ed.get("entry", ""))
        ref_type = html.escape(ed.get("type", "other"))
        all_issues: list[str] = ed.get("all_issues", [])
        suggestion = html.escape(ed.get("suggestion", ""))
        alpha_ok = ed.get("alphabetical_ok", True)
        needs_manual = ed.get("needs_manual_lookup", False)

        # Status badge
        if not all_issues and not needs_manual:
            status_html = '<span class="badge badge-ok">OK</span>'
        elif needs_manual:
            status_html = '<span class="badge badge-warn">Check</span>'
        else:
            status_html = '<span class="badge badge-error">Issues</span>'

        # Alpha badge
        if alpha_ok:
            alpha_badge = '<span class="alpha-badge-ok">✓</span>'
        else:
            alpha_badge = '<span class="alpha-badge-bad">✗</span>'

        # Issues list
        if all_issues:
            issue_items = '\n'.join(f'<li>{html.escape(i)}</li>' for i in all_issues)
            issues_html = f'<ul class="issue-list">{issue_items}</ul>'
        else:
            issues_html = '<span class="badge badge-ok">None</span>'

        # Suggestion
        sugg_html = f'<div class="sugg">{suggestion}</div>' if suggestion else ''

        rows.append(f"""
<tr>
  <td>{idx}</td>
  <td><span class="orig">{entry_text}</span></td>
  <td><span class="type-tag">{ref_type}</span></td>
  <td>{alpha_badge}</td>
  <td>{status_html}</td>
  <td>{issues_html}</td>
  <td>{sugg_html}</td>
</tr>""")

    table_body = '\n'.join(rows)

    # ---------- Assemble ----------
    html_content = f"""<!DOCTYPE html>
<html lang="pt">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>APA 7 Reference Check Report</title>
  <style>
{_HTML_CSS}
  </style>
</head>
<body>
<div class="container">
  <h1>APA 7 Reference Check Report</h1>
  <p class="meta">Source: {source_file}</p>
  <p class="meta">Generated: {timestamp}</p>

  <h2>Summary</h2>
  {cards_html}

  <h2>Alphabetical Order Check</h2>
  {alpha_html}

  <h2>Issue Breakdown</h2>
  <table>
    <thead>
      <tr>
        <th>#</th>
        <th>Entry</th>
        <th>Type</th>
        <th>Order</th>
        <th>Status</th>
        <th>Issues</th>
        <th>Suggestion</th>
      </tr>
    </thead>
    <tbody>
      {table_body}
    </tbody>
  </table>
</div>
</body>
</html>"""

    if out_path:
        safe_write_file(out_path, html_content)

    return html_content


# ---------------------------------------------------------------------------
# 26. safe_write_file
# ---------------------------------------------------------------------------

def safe_write_file(
    path: "str | Path",
    content: str,
    encoding: str = "utf-8",
) -> str:
    """Write *content* to *path* with a 4-level fallback strategy.

    1. Try original path.
    2. Try accent-stripped version.
    3. Try ``%TEMP%/apa7_output/{filename}``.
    4. Try system temp dir.

    Returns the path where the file was actually written.
    """
    path = Path(path)

    def _try_write(p: Path) -> bool:
        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            p.write_text(content, encoding=encoding)
            return True
        except (OSError, UnicodeEncodeError):
            return False

    # 1. Original
    if _try_write(path):
        return str(path)

    # 2. Accent-stripped
    stripped_name = _strip_accents(path.name)
    stripped_path = path.parent / stripped_name
    if _try_write(stripped_path):
        return str(stripped_path)

    # 3. %TEMP%/apa7_output/
    temp_dir = Path(tempfile.gettempdir()) / 'apa7_output'
    temp_path = temp_dir / _strip_accents(path.name)
    if _try_write(temp_path):
        return str(temp_path)

    # 4. System temp dir directly
    last_resort = Path(tempfile.gettempdir()) / _strip_accents(path.name)
    if _try_write(last_resort):
        return str(last_resort)

    raise OSError(f"Could not write file to any fallback location for: {path}")


# ---------------------------------------------------------------------------
# 27. check_docx_references — full pipeline
# ---------------------------------------------------------------------------

def check_docx_references(
    docx_path: "str | Path",
    rules_path: "str | Path | None" = None,
    out_json: "str | Path | None" = None,
    out_html: "str | Path | None" = None,
    out_word_instructions: "str | Path | None" = None,
) -> dict:
    """Full APA 7 reference-checking pipeline for a .docx file.

    Steps:
      1. Load rules (optional).
      2. Extract full document text.
      3. Find the references section.
      4. Split into individual entries.
      5. Extract italic formatting spans.
      6. Validate each entry (heuristic + type-specific + italic).
      7. Check alphabetical order.
      8. Build report dict.
      9. Write outputs (JSON, HTML, Word instructions).

    Returns the full report dict.
    """
    docx_path = Path(docx_path)

    # 1. Load rules
    rules: dict = {}
    if rules_path:
        rules = load_apa_rules(rules_path)

    # 2. Extract text
    full_text = extract_docx_text(docx_path)

    # 3. Find references section
    refs_text = find_references_section(full_text)

    # 4. Split references
    entries = split_references(refs_text)

    # 5. Italic lookup
    try:
        italic_lookup = extract_docx_refs_with_italic(docx_path)
    except Exception:
        italic_lookup = {}

    # 6. Validate each entry
    entry_results: list[dict] = []
    for entry in entries:
        # Heuristic validation
        heuristic = validate_reference_entry(entry, rules)
        issues: list[str] = list(heuristic.get("issues", []))
        suggestion = heuristic.get("suggestion", "")

        # Type classification
        ref_type = classify_reference_type(entry, rules)

        # Type-specific validation
        type_result = validate_by_type(entry, ref_type)
        type_issues: list[str] = type_result.get("type_issues", [])
        issues.extend(type_issues)
        if type_result.get("suggestion") and not suggestion:
            suggestion = type_result["suggestion"]

        # Jammed entries
        jammed = detect_jammed_entries(entry)
        if jammed:
            issues.append(
                f"Possible jammed entry: may contain {len(jammed)} merged references"
            )

        # Italic formatting
        italic_data = italic_lookup.get(entry.strip(), {})
        italic_spans = italic_data.get("italic_spans", [])
        italic_issues = validate_italic_formatting(entry, italic_spans, ref_type)
        issues.extend(italic_issues)

        needs_manual = bool(issues) or bool(jammed)

        entry_results.append({
            "entry": entry,
            "type": ref_type,
            "all_issues": issues,
            "suggestion": suggestion,
            "jammed_split": jammed,
            "needs_manual_lookup": needs_manual,
            "alphabetical_ok": True,  # filled in step 8
        })

    # 7. Alphabetical order
    alpha_result = check_alphabetical_order(entries)
    out_of_order_set = set(alpha_result.get("out_of_order", []))

    # 8. Tag each entry
    for er in entry_results:
        er["alphabetical_ok"] = er["entry"] not in out_of_order_set

    # Build report
    timestamp = datetime.now().isoformat(timespec='seconds')
    report: dict = {
        "source_file": str(docx_path),
        "timestamp": timestamp,
        "total_references": len(entries),
        "entries": entry_results,
        "alphabetical_order": alpha_result,
    }

    # 9. Write JSON
    if out_json:
        json_str = json.dumps(report, ensure_ascii=False, indent=2)
        safe_write_file(out_json, json_str)

    # 10. Write HTML
    if out_html:
        generate_html_report(report, out_path=out_html)

    # 11. Write Word instructions
    if out_word_instructions:
        word_instr = alpha_result.get("word_instructions", "")
        safe_write_file(out_word_instructions, word_instr)

    return report
